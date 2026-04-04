import os
import json
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from app.services.word_service import (
    load_template,
    generate_word_document,
    _add_paragraph_styled,
    _add_numbered_question,
    _add_multiple_choice,
    _add_jawaban,
    _add_isian_essay,
)


def create_test_template(tmp_path: Path) -> str:
    from docx import Document
    doc = Document()
    doc.add_heading("{{JUDUL_UJIAN}}", level=1)
    doc.add_paragraph("Nama: {{NAMA_SISWA}}")
    doc.add_paragraph("Kelas: {{KELAS}}")
    doc.add_paragraph("Tanggal: {{TANGGAL}}")
    template_path = tmp_path / "test_template.docx"
    doc.save(str(template_path))
    return str(template_path)


class TestLoadTemplate:
    def test_load_existing_template(self, tmp_path):
        template_path = create_test_template(tmp_path)
        doc = load_template(template_path)
        assert doc is not None
        assert len(doc.paragraphs) > 0

    def test_load_nonexistent_template_raises_error(self):
        with pytest.raises(FileNotFoundError, match="Template tidak ditemukan"):
            load_template("/nonexistent/path/template.docx")


class TestAddParagraphStyled:
    def test_add_bold_paragraph(self, tmp_path):
        from docx import Document
        doc = Document()
        para = _add_paragraph_styled(doc, "Test Text", bold=True, size=12)
        assert len(doc.paragraphs) == 1
        assert doc.paragraphs[0].runs[0].bold is True
        assert doc.paragraphs[0].runs[0].text == "Test Text"

    def test_add_normal_paragraph(self, tmp_path):
        from docx import Document
        doc = Document()
        para = _add_paragraph_styled(doc, "Normal Text", bold=False, size=11)
        assert doc.paragraphs[0].runs[0].bold is False


class TestAddNumberedQuestion:
    def test_add_numbered_question(self, tmp_path):
        from docx import Document
        doc = Document()
        _add_numbered_question(doc, 1, "Apa itu Python?")
        assert "1. Apa itu Python?" in doc.paragraphs[0].runs[0].text
        assert doc.paragraphs[0].runs[0].bold is True


class TestAddMultipleChoice:
    def test_add_multiple_choice(self, tmp_path):
        from docx import Document
        doc = Document()
        pilihan = ["A. Jawaban 1", "B. Jawaban 2", "C. Jawaban 3", "D. Jawaban 4"]
        _add_multiple_choice(doc, pilihan)
        assert len(doc.paragraphs) == 4
        assert "A. Jawaban 1" in doc.paragraphs[0].runs[0].text


class TestAddJawaban:
    def test_add_jawaban_with_pembahasan(self, tmp_path):
        from docx import Document
        doc = Document()
        _add_jawaban(doc, "B", "Penjelasan lengkap")
        assert len(doc.paragraphs) == 2
        assert "Kunci Jawaban: B" in doc.paragraphs[0].runs[0].text
        assert "Pembahasan: Penjelasan lengkap" in doc.paragraphs[1].runs[0].text

    def test_add_jawaban_without_pembahasan(self, tmp_path):
        from docx import Document
        doc = Document()
        _add_jawaban(doc, "A")
        assert len(doc.paragraphs) == 1
        assert "Kunci Jawaban: A" in doc.paragraphs[0].runs[0].text


class TestAddIsianEssay:
    def test_add_isian_essay(self, tmp_path):
        from docx import Document
        doc = Document()
        _add_isian_essay(doc, 1, "Jelaskan konsep ini", "Jawaban benar", "Penjelasan")
        assert len(doc.paragraphs) == 4
        assert "1. Jelaskan konsep ini" in doc.paragraphs[0].runs[0].text
        assert "Jawaban: ___" in doc.paragraphs[1].runs[0].text


class TestGenerateWordDocument:
    def test_generate_document_dari_template(self, tmp_path):
        template_path = create_test_template(tmp_path)
        soal_data = [
            {
                "nomor": 1,
                "pertanyaan": "Berapa 2 + 2?",
                "pilihan": ["A. 3", "B. 4", "C. 5", "D. 6"],
                "jawaban": "B",
                "pembahasan": "2 + 2 = 4",
            }
        ]

        output_path = generate_word_document(
            soal_data=soal_data,
            template_path=template_path,
            judul_ujian="Matematika Dasar",
            nama_siswa="Test Siswa",
            kelas="7A",
            tanggal="2026-04-04",
        )

        assert os.path.exists(output_path)
        assert output_path.endswith(".docx")

        doc = load_template(output_path)
        assert len(doc.paragraphs) > 0

    def test_generate_dengan_kunci_terpisah(self, tmp_path):
        template_path = create_test_template(tmp_path)
        soal_data = [
            {
                "nomor": 1,
                "pertanyaan": "Test?",
                "pilihan": ["A. 1", "B. 2"],
                "jawaban": "A",
                "pembahasan": "Test pembahasan",
            },
            {
                "nomor": 2,
                "pertanyaan": "Jelaskan!",
                "jawaban": "Jawaban essay",
                "pembahasan": "Pembahasan essay",
            },
        ]

        output_path = generate_word_document(
            soal_data=soal_data,
            template_path=template_path,
            judul_ujian="Ujian Akhir",
            sertakan_kunci_terpisah=True,
        )

        assert os.path.exists(output_path)
        doc = load_template(output_path)
        texts = [p.text for p in doc.paragraphs]
        assert any("KUNCI JAWABAN" in t for t in texts)

    def test_generate_placeholder_replacement(self, tmp_path):
        template_path = create_test_template(tmp_path)
        soal_data = [
            {
                "nomor": 1,
                "pertanyaan": "Test?",
                "pilihan": ["A. 1"],
                "jawaban": "A",
            }
        ]

        output_path = generate_word_document(
            soal_data=soal_data,
            template_path=template_path,
            judul_ujian="Matematika",
            nama_siswa="Budi",
            kelas="7A",
            tanggal="2026-04-04",
        )

        doc = load_template(output_path)
        all_text = " ".join([p.text for p in doc.paragraphs])
        assert "{{JUDUL_UJIAN}}" not in all_text or "Matematika" in all_text
