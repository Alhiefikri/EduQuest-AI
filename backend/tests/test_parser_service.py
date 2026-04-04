import io
import pytest

from app.services.parser_service import extract_text_from_pdf, extract_text_from_docx, count_words


def create_test_pdf(text: str = "Hello World from PDF") -> bytes:
    import fitz
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((50, 50), text)
    return doc.tobytes()


def create_test_docx(text: str = "Hello World from Word") -> bytes:
    from docx import Document
    doc = Document()
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


class TestExtractTextFromPDF:
    def test_extract_text_from_valid_pdf(self):
        pdf_content = create_test_pdf("Modul Ajar Matematika Kelas 7")
        text, page_count = extract_text_from_pdf(pdf_content)
        assert text is not None
        assert len(text) > 0
        assert page_count >= 1
        assert "Matematika" in text

    def test_extract_text_from_empty_pdf_raises_error(self):
        invalid_pdf = b"not a valid pdf"
        with pytest.raises(ValueError, match="rusak atau tidak dapat dibaca"):
            extract_text_from_pdf(invalid_pdf)


class TestExtractTextFromDocx:
    def test_extract_text_from_valid_docx(self):
        docx_content = create_test_docx("Modul Ajar Bahasa Indonesia")
        text, page_count = extract_text_from_docx(docx_content)
        assert text is not None
        assert len(text) > 0
        assert "Bahasa Indonesia" in text
        assert page_count >= 1

    def test_extract_text_from_docx_with_tables(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("Tabel Data")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Nama"
        table.cell(0, 1).text = "Nilai"
        table.cell(1, 0).text = "Andi"
        table.cell(1, 1).text = "90"
        buffer = io.BytesIO()
        doc.save(buffer)

        text, page_count = extract_text_from_docx(buffer.getvalue())
        assert "Tabel Data" in text
        assert "Andi" in text
        assert "90" in text

    def test_extract_text_from_empty_docx_raises_error(self):
        from docx import Document
        doc = Document()
        buffer = io.BytesIO()
        doc.save(buffer)
        with pytest.raises(ValueError, match="tidak mengandung teks"):
            extract_text_from_docx(buffer.getvalue())

    def test_extract_text_from_corrupt_docx_raises_error(self):
        invalid_docx = b"not a valid docx"
        with pytest.raises(ValueError, match="rusak atau tidak dapat dibaca"):
            extract_text_from_docx(invalid_docx)


class TestCountWords:
    def test_count_words_normal(self):
        assert count_words("Hello World") == 2
        assert count_words("Modul Ajar Matematika Kelas 7") == 5

    def test_count_words_empty_string(self):
        assert count_words("") == 0

    def test_count_words_single_word(self):
        assert count_words("Hello") == 1

    def test_count_words_with_extra_spaces(self):
        assert count_words("  Hello   World  ") == 2
