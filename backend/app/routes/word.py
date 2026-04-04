import json
import os
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from fastapi import APIRouter, UploadFile, File, HTTPException, Query
from fastapi.responses import FileResponse

from app.config import TEMPLATES_DIR, MAX_FILE_SIZE
from app.database.connection import get_db
from app.models.word import (
    TemplateUploadResponse,
    TemplateListResponse,
    GenerateWordRequest,
    GenerateWordResponse,
)
from app.services.word_service import generate_word_document

router = APIRouter(prefix="/api/v1/word", tags=["word"])


def _create_default_template() -> str:
    default_path = TEMPLATES_DIR / "default.docx"
    if not default_path.exists():
        doc = DocxDocument()
        doc.add_heading("{{JUDUL_UJIAN}}", level=1)
        doc.add_paragraph("Nama: {{NAMA_SISWA}}")
        doc.add_paragraph("Kelas: {{KELAS}}")
        doc.add_paragraph("Tanggal: {{TANGGAL}}")
        doc.save(str(default_path))
    return str(default_path)


@router.post("/template", response_model=TemplateUploadResponse, status_code=201)
async def upload_template(file: UploadFile = File(...), nama: str = Query(None)):
    if file.content_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        raise HTTPException(status_code=400, detail="Hanya file Word (.docx) yang diterima")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="Ukuran file melebihi batas maksimum 10MB")

    file_ext = ".docx"
    unique_filename = f"{uuid.uuid4().hex}{file_ext}"
    file_path = TEMPLATES_DIR / unique_filename

    with open(file_path, "wb") as f:
        f.write(content)

    template_name = nama or file.filename or "Template Tanpa Nama"

    db = await get_db()

    try:
        template = await db.templateWord.create(
            data={
                "nama": template_name,
                "filePath": str(file_path),
                "isDefault": False,
            }
        )
    except Exception as e:
        if file_path.exists():
            os.remove(file_path)
        raise HTTPException(status_code=500, detail=f"Gagal menyimpan template ke basis data: {str(e)}")

    return TemplateUploadResponse.from_prisma(template)


@router.get("/template", response_model=list[TemplateListResponse])
async def list_templates():
    db = await get_db()

    try:
        templates = await db.templateWord.find_many(order={"createdAt": "desc"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengambil daftar template: {str(e)}")

    return [TemplateListResponse.from_prisma(t) for t in templates]


@router.get("/template/{template_id}", response_model=TemplateUploadResponse)
async def get_template(template_id: str):
    db = await get_db()

    try:
        template = await db.templateWord.find_unique(where={"id": template_id})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengambil template: {str(e)}")

    if not template:
        raise HTTPException(status_code=404, detail="Template tidak ditemukan")

    return TemplateUploadResponse.from_prisma(template)


@router.post("/generate", response_model=GenerateWordResponse)
async def generate_word(request: GenerateWordRequest):
    db = await get_db()

    try:
        soal = await db.soal.find_unique(where={"id": request.soal_id})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengambil data soal: {str(e)}")

    if not soal:
        raise HTTPException(status_code=404, detail="Soal tidak ditemukan")

    soal_data = json.loads(soal.dataSoal)

    template_path = None
    if request.template_id:
        try:
            template = await db.templateWord.find_unique(where={"id": request.template_id})
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Gagal mengambil template: {str(e)}")

        if not template:
            raise HTTPException(status_code=404, detail="Template tidak ditemukan")

        template_path = template.filePath
    else:
        default_template = await db.templateWord.find_first(where={"isDefault": True})
        if default_template:
            template_path = default_template.filePath
        else:
            template_path = _create_default_template()

    try:
        output_path = generate_word_document(
            soal_data=soal_data,
            template_path=template_path,
            judul_ujian=request.judul_ujian or f"Soal {soal.mataPelajaran}",
            nama_siswa=request.nama_siswa or "________________",
            kelas=request.kelas or "________________",
            tanggal=request.tanggal or "________________",
            sertakan_kunci_terpisah=request.sertakan_kunci_terpisah,
        )
    except FileNotFoundError as e:
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal membuat dokumen Word: {str(e)}")

    output_filename = Path(output_path).name

    try:
        await db.soal.update(
            where={"id": request.soal_id},
            data={"filePath": output_path},
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Dokumen Word berhasil dibuat, namun gagal mencatat path file ke basis data: {str(e)}",
        )

    return GenerateWordResponse(
        file_path=output_path,
        file_name=output_filename,
        message="Dokumen Word berhasil dibuat",
    )


@router.get("/download/{soal_id}")
async def download_word(soal_id: str):
    db = await get_db()

    try:
        soal = await db.soal.find_unique(where={"id": soal_id})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Gagal mengambil data soal: {str(e)}")

    if not soal:
        raise HTTPException(status_code=404, detail="Soal tidak ditemukan")

    if not soal.filePath:
        raise HTTPException(status_code=404, detail="File Word belum di-generate. Silakan generate terlebih dahulu.")

    file_path = Path(soal.filePath)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File Word tidak ditemukan di penyimpanan")

    return FileResponse(
        path=str(file_path),
        filename=file_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
