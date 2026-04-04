import io
from typing import List, Tuple

import fitz  # type: ignore
from docx import Document as DocxDocument  # type: ignore


def extract_text_from_pdf(file_content: bytes) -> Tuple[str, int]:
    try:
        doc = fitz.open(stream=file_content, filetype="pdf")
    except Exception as e:
        raise ValueError(f"File PDF rusak atau tidak dapat dibaca: {str(e)}")

    text_parts = []
    page_count = doc.page_count

    for page_num in range(page_count):
        page = doc[page_num]
        page_text = page.get_text()
        text_parts.append(page_text)

    doc.close()

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError("File PDF tidak mengandung teks. Kemungkinan file hanya berisi gambar.")

    return full_text, page_count


def extract_text_from_docx(file_content: bytes) -> Tuple[str, int]:
    try:
        doc = DocxDocument(io.BytesIO(file_content))
    except Exception as e:
        raise ValueError(f"File Word rusak atau tidak dapat dibaca: {str(e)}")

    text_parts = []

    for para in doc.paragraphs:
        text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text_parts.append(cell.text)

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError("File Word tidak mengandung teks")

    # Estimate: ~3000 characters per page (average A4 font 12pt)
    page_count = max(1, len(full_text) // 3000)

    return full_text, page_count


def extract_text_from_pdf_by_pages(file_path: str, page_numbers: List[int]) -> str:
    """
    Ekstrak teks dari file PDF pada halaman-halaman tertentu.
    page_numbers: list of 1-based page numbers.
    """
    try:
        doc = fitz.open(file_path)
    except Exception as e:
        raise ValueError(f"Gagal membuka file PDF dari disk: {str(e)}")

    text_parts = []
    max_pages = doc.page_count

    for p_num in page_numbers:
        # Convert 1-based to 0-based
        idx = p_num - 1
        if 0 <= idx < max_pages:
            page = doc[idx]
            text_parts.append(page.get_text())

    doc.close()
    return "\n".join(text_parts).strip()


def count_words(text: str) -> int:
    return len(text.split())
