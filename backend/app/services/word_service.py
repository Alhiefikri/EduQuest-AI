import json
import os
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.config import TEMPLATES_DIR, OUTPUTS_DIR


def load_template(template_path: str) -> DocxDocument:
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template tidak ditemukan: {template_path}")
    return DocxDocument(template_path)


def _add_paragraph_styled(doc: DocxDocument, text: str, bold: bool = False, size: int = 11, alignment=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if alignment:
        para.alignment = alignment
    return para


def _add_numbered_question(doc: DocxDocument, nomor: int, pertanyaan: str):
    _add_paragraph_styled(doc, f"{nomor}. {pertanyaan}", bold=True, size=12)


def _add_multiple_choice(doc: DocxDocument, pilihan: list):
    for opsi in pilihan:
        _add_paragraph_styled(doc, f"   {opsi}", size=11)


def _add_jawaban(doc: DocxDocument, jawaban: str, pembahasan: str = None):
    _add_paragraph_styled(doc, f"Kunci Jawaban: {jawaban}", bold=True, size=11)
    if pembahasan:
        _add_paragraph_styled(doc, f"Pembahasan: {pembahasan}", size=11)


def _add_isian_essay(doc: DocxDocument, nomor: int, pertanyaan: str, jawaban: str = None, pembahasan: str = None):
    _add_paragraph_styled(doc, f"{nomor}. {pertanyaan}", bold=True, size=12)
    _add_paragraph_styled(doc, "Jawaban: ________________________________", size=11)
    if jawaban:
        _add_paragraph_styled(doc, f"Kunci Jawaban: {jawaban}", bold=True, size=11)
    if pembahasan:
        _add_paragraph_styled(doc, f"Pembahasan: {pembahasan}", size=11)


def generate_word_document(
    soal_data: list,
    template_path: str,
    judul_ujian: str = "Ujian",
    nama_siswa: str = "________________",
    kelas: str = "________________",
    tanggal: str = "________________",
    sertakan_kunci_terpisah: bool = False,
) -> str:
    doc = load_template(template_path)

    for placeholder, value in [
        ("{{JUDUL_UJIAN}}", judul_ujian),
        ("{{NAMA_SISWA}}", nama_siswa),
        ("{{KELAS}}", kelas),
        ("{{TANGGAL}}", tanggal),
    ]:
        for para in doc.paragraphs:
            if placeholder in para.text:
                for run in para.runs:
                    run.text = run.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(placeholder, value)

    doc.add_page_break()
    _add_paragraph_styled(doc, judul_ujian, bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    for item in soal_data:
        nomor = item.get("nomor", 1)
        pertanyaan = item.get("pertanyaan", "")
        pilihan = item.get("pilihan", [])
        jawaban = item.get("jawaban", "")
        pembahasan = item.get("pembahasan")
        tipe = "pilihan_ganda" if pilihan else "isian"

        if tipe == "pilihan_ganda":
            _add_numbered_question(doc, nomor, pertanyaan)
            _add_multiple_choice(doc, pilihan)
            if sertakan_kunci_terpisah:
                pass
            else:
                _add_jawaban(doc, jawaban, pembahasan)
        else:
            _add_isian_essay(doc, nomor, pertanyaan, jawaban if not sertakan_kunci_terpisah else None, pembahasan if not sertakan_kunci_terpisah else None)

        doc.add_paragraph()

    if sertakan_kunci_terpisah:
        doc.add_page_break()
        _add_paragraph_styled(doc, "KUNCI JAWABAN", bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        doc.add_paragraph()
        for item in soal_data:
            nomor = item.get("nomor", 1)
            jawaban = item.get("jawaban", "")
            pembahasan = item.get("pembahasan")
            _add_paragraph_styled(doc, f"{nomor}. {jawaban}", bold=True, size=11)
            if pembahasan:
                _add_paragraph_styled(doc, f"   Pembahasan: {pembahasan}", size=11)

    OUTPUTS_DIR.mkdir(parents=True, exist_ok=True)
    output_filename = f"soal_{uuid.uuid4().hex[:8]}.docx"
    output_path = OUTPUTS_DIR / output_filename

    doc.save(str(output_path))

    return str(output_path)
